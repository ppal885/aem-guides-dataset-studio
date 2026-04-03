"""
Generate Prashant Pal's resume as DOCX with AEM Guides Dataset Studio project.
Run: python scripts/generate_resume_docx.py
"""
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(12)
    p.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    # Header
    p = doc.add_paragraph()
    run = p.add_run("PRASHANT PAL")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("SDET | AI-Powered Test Automation | Selenium Python BDD Expert")
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run("Noida, India | palsintu3@gmail.com | github.com/prashantp")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(12)

    # Professional Summary
    add_heading(doc, "PROFESSIONAL SUMMARY", 1)
    doc.add_paragraph(
        "Results-driven Senior SDET with 5+ years of experience in designing, building, and maintaining "
        "enterprise-grade test automation frameworks. Specialized in Selenium WebDriver, Python, and BDD "
        "(Behave/Cucumber) for large-scale SaaS products. Pioneered an AI-powered agentic test automation pipeline "
        "that integrates Jira ticket ingestion, automated test generation, self-healing execution, flakiness analysis, "
        "and Allure reporting — reducing test creation time by 70% and improving test stability across 3,500+ "
        "automated scenarios. Built AEM Guides Dataset Studio — an AI pipeline for DITA dataset generation from Jira "
        "with RAG, multi-stage LLM classification, and self-learning feedback loops. Proven track record of building "
        "Page Object Model architectures, CI/CD test pipelines, and cross-browser test infrastructure for Adobe "
        "Experience Manager products."
    )
    doc.add_paragraph()

    # Core Competencies
    add_heading(doc, "CORE COMPETENCIES / TECHNICAL SKILLS", 1)
    add_bullet(doc, "Test Automation: Selenium WebDriver, Python, Behave (BDD/Cucumber), Page Object Model (POM), XPath Selectors, Cross-Browser Testing (Chrome, Firefox), Headless Testing, Parallel Execution")
    add_bullet(doc, "AI/ML in Testing: AI-Powered Test Generation, Self-Healing Test Framework, Agentic Automation Pipeline, LLM-Driven Code Generation, RAG Pipelines (ChromaDB, LangChain), Multi-Stage LLM Classification, Self-Learning Feedback Loops, Browser MCP (Model Context Protocol) for UI Exploration")
    add_bullet(doc, "Frameworks & Patterns: Custom Element/Page/Widget Architecture, Polling Conditions, Retry Mechanisms, Dynamic XPath Templates, Scrolling/Viewport Handling")
    add_bullet(doc, "CI/CD & DevOps: Jenkins, GitHub Actions, Docker, Allure Reporting, PostgreSQL Test Tracking, Headless Chrome Automation, Parallel Test Execution")
    add_bullet(doc, "Programming: Python (Expert), JavaScript, SQL, Bash/PowerShell, REST API Testing")
    add_bullet(doc, "Tools & Platforms: Jira, Git, Adobe Experience Manager (AEM), Allure, VS Code/Cursor IDE, npm, pip, PostgreSQL, Spectrum Design System")
    add_bullet(doc, "Testing Types: Functional, Regression, Integration, End-to-End, Smoke/BVT, Performance Benchmarking, Flakiness Analysis, Cross-Platform (Windows/Linux)")
    doc.add_paragraph()

    # Professional Experience
    add_heading(doc, "PROFESSIONAL EXPERIENCE", 1)

    add_heading(doc, "SDET | Adobe Systems", 2)
    doc.add_paragraph("Noida, India | 2022 – Present")
    doc.add_paragraph()

    add_bullet(doc, "Architected and maintained an enterprise Selenium Python BDD (Behave) test automation framework covering 3,500+ step definitions, 60+ page objects, 90+ widget components, and 1,000+ XPath selectors for Adobe Experience Manager Guides — a mission-critical SaaS content management platform.")
    add_bullet(doc, "Pioneered an AI-powered agentic test automation pipeline integrating 8 specialized AI agents (orchestrator, planner, generator, runner, healer, step-implementer, feature-writer, flakiness-reporter) that automate the full lifecycle from Jira ticket to Allure report with self-healing capabilities.")
    add_bullet(doc, "Built AEM Guides Dataset Studio — full-stack AI pipeline (FastAPI + React) that generates DITA datasets from Jira evidence using multi-stage LLM classification (mechanism → pattern → recipe routing), RAG (ChromaDB, DITA spec, AEM Guides docs), and 30+ deterministic recipes with LLM fallback for novel constructs.")
    add_bullet(doc, "Implemented self-learning feedback loop in Dataset Studio: routing and prompt overrides from user corrections applied automatically without retraining; ChatGPT-style paste flow with tool-calling, streaming, and RAG grounding.")
    add_bullet(doc, "Built a self-learning knowledge base with an auto-generated framework index (3,517 steps, 201 utilities, 38 polling conditions), selector registry for XPath migration tracking, fix-pattern log, and timing benchmarks — enabling AI agents to reuse existing code and avoid duplication.")
    add_bullet(doc, "Designed and implemented a flakiness analysis system using PostgreSQL stability metrics (consecutive_fail_count, stability_pass/fail_count) and Allure result parsing to generate actionable flakiness scores, timing regression detection (>20% threshold), and per-owner reports.")
    add_bullet(doc, "Reduced test creation time by 70% through AI-driven code generation that searches the framework index before producing new code, enforcing mandatory reuse checks across 6 data sources.")
    add_bullet(doc, "Implemented browser-based UI exploration via Model Context Protocol (MCP) for automated XPath discovery, live DOM inspection, and selector validation — replacing manual element inspection.")
    add_bullet(doc, "Established strict code quality standards enforced via Cursor rules: no unnecessary logging, no try/catch in step definitions, no hard sleeps (using Visible/Clickable/PollerCondition waits), and all static XPaths initialized in page object constructors.")
    add_bullet(doc, "Engineered cross-environment test execution supporting Cloud UUID, On-Prem UUID, and On-Prem Non-UUID AEM deployments with dynamic tag calculation, headless Chrome (3072x1920 viewport), and parallel execution.")
    add_bullet(doc, "Built comprehensive Allure reporting infrastructure with failure screenshots, step-level timing, and CI/CD integration producing reports for UUID, Non-UUID, Cloud, and BVT test suites.")
    add_bullet(doc, "Maintained 95%+ test stability across 3,500+ scenarios through systematic flakiness reduction, selector migration tracking, and automated self-healing of broken tests.")
    doc.add_paragraph()

    add_heading(doc, "QA Automation Engineer | Previous Company", 2)
    doc.add_paragraph("India | 2017 – 2019")
    doc.add_paragraph()
    add_bullet(doc, "Developed automated test suites using Selenium WebDriver with Python for web application testing, achieving 80% automation coverage across critical user flows.")
    add_bullet(doc, "Implemented BDD test framework using Behave/Cucumber with Gherkin feature files, enabling collaboration between QA, developers, and product managers on test specifications.")
    add_bullet(doc, "Created and maintained Page Object Model architecture for 20+ page objects covering authentication, dashboards, forms, and reporting modules.")
    add_bullet(doc, "Performed REST API testing using Python requests library with JSON schema validation and data-driven parameterization for comprehensive endpoint coverage.")
    add_bullet(doc, "Collaborated with development teams in Agile/Scrum methodology, participating in sprint planning, daily standups, and retrospectives to ensure quality throughout the SDLC.")
    doc.add_paragraph()

    # Key Projects
    add_heading(doc, "KEY PROJECTS", 1)

    add_heading(doc, "AI-Powered Agentic Test Automation Pipeline", 2)
    add_bullet(doc, "End-to-end pipeline: Jira ticket ingestion → AI test plan generation → code generation → execution with Allure → self-healing → flakiness reporting — fully automated in a single command.")
    add_bullet(doc, "8 specialized AI agents orchestrated via Cursor IDE with MCP integrations for Jira and browser automation.")
    add_bullet(doc, "Self-learning knowledge base: auto-indexed 3,517 step definitions, 1,081 XPath selectors, and 201 utility functions to prevent code duplication and enforce framework pattern reuse.")
    add_bullet(doc, "Flakiness scorer computing stability metrics from PostgreSQL + Allure data, categorizing scenarios as Stable/Flaky/Broken with owner-level accountability reports.")
    add_bullet(doc, "Timing benchmark system detecting performance regressions >20% with headless vs headed mode tracking.")
    doc.add_paragraph()

    add_heading(doc, "AEM Guides Dataset Studio — AI-Powered DITA Dataset Generator", 2)
    add_bullet(doc, "Full-stack AI pipeline (FastAPI + React) that generates DITA datasets from Jira evidence using multi-stage LLM classification (mechanism → pattern → recipe routing), RAG (ChromaDB, DITA spec, AEM Guides docs), and 30+ deterministic recipes with LLM fallback for novel constructs.")
    add_bullet(doc, "Self-learning feedback loop: routing overrides from wrong-recipe corrections, prompt overrides (deprioritize/prefer recipes) applied automatically without retraining.")
    add_bullet(doc, "ChatGPT-style paste flow with tool-calling, streaming, and RAG grounding; evaluation framework for domain accuracy and validation rate.")
    add_bullet(doc, "Tech: Python, FastAPI, Claude/Groq, LangChain, ChromaDB, React, Vite, TypeScript, SQLite")
    doc.add_paragraph()

    # Certifications
    add_heading(doc, "CERTIFICATIONS", 1)
    add_bullet(doc, "ISTQB Certified Tester — Foundation Level")
    add_bullet(doc, "Selenium WebDriver with Python — Advanced Automation")
    add_bullet(doc, "AWS Cloud Practitioner (or relevant cloud certification)")
    doc.add_paragraph()

    # Education
    add_heading(doc, "EDUCATION", 1)
    doc.add_paragraph("Bachelor of Engineering / Bachelor of Technology in Computer Science")
    doc.add_paragraph("University Name, India | Graduated: 20XX")
    doc.add_paragraph()

    # Achievements
    add_heading(doc, "ACHIEVEMENTS & RECOGNITION", 1)
    add_bullet(doc, "Reduced test creation cycle from 2 days to 4 hours through AI-powered test generation pipeline.")
    add_bullet(doc, "Achieved 95%+ test stability across 3,500+ automated scenarios in enterprise SaaS product.")
    add_bullet(doc, "Indexed and cataloged entire test framework (3,517 steps, 60+ pages, 90+ widgets) for AI-driven reuse.")
    add_bullet(doc, "Eliminated 40% code duplication through Widget pattern and reusable component architecture.")
    add_bullet(doc, "Designed flakiness detection system adopted by the QA team for data-driven test maintenance prioritization.")
    add_bullet(doc, "Built AEM Guides Dataset Studio — full-stack AI pipeline for DITA generation with RAG, self-learning, and ChatGPT-style UI.")

    # Save
    out_dir = Path(__file__).resolve().parent.parent
    out_path = out_dir / "Prashant_Pal_Resume.docx"
    doc.save(str(out_path))
    print(f"Resume saved to: {out_path}")
    print(f"Full path: {out_path.resolve()}")
    return str(out_path)


if __name__ == "__main__":
    main()
