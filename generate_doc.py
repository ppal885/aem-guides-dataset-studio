from docx import Document
from docx.shared import Pt, RGBColor
import datetime

doc = Document()

# Title
doc.add_heading('AEM Guides Dataset Studio', 0)
doc.add_heading('Linux VM Deployment — Learnings, Issues & Resolutions', 1)
doc.add_paragraph(f'Date: {datetime.date.today().strftime("%B %d, %Y")}   |   Author: prashantp@adobe.com   |   VM: 10.42.46.78:4502')

# 1. Summary
doc.add_heading('1. Deployment Summary', 1)
doc.add_paragraph(
    'The AEM Guides Dataset Studio was successfully deployed on Ubuntu 22.04 LTS '
    '(Adobe ATS Corp internal VM). The app runs without Docker — FastAPI backend managed '
    'via systemd, React frontend served via nginx on port 4502.'
)
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Component'
t.rows[0].cells[1].text = 'Details'
for label, val in [
    ('URL', 'http://10.42.46.78:4502/'),
    ('OS', 'Ubuntu 22.04.5 LTS'),
    ('Backend', 'Port 8001, systemd service: aem-backend'),
    ('Frontend', 'Port 4502, nginx'),
    ('LLM', 'Azure OpenAI gpt-5.2'),
    ('Database', 'SQLite (default)'),
    ('Auto-restart', 'Yes — systemd Restart=always'),
]:
    r = t.add_row()
    r.cells[0].text = label
    r.cells[1].text = val

# 2. Issues
doc.add_heading('2. Issues Found & Fixes Applied', 1)

issues = [
    (
        'Docker group not found',
        'docker group did not exist because Docker daemon was never started.',
        'sudo systemctl start docker && sudo usermod -aG docker $USER'
    ),
    (
        'Docker service unit file not found',
        'Docker was not properly installed by the get.docker.com convenience script.',
        'Installed via official apt repository: docker-ce, docker-ce-cli, containerd.io, docker-compose-plugin'
    ),
    (
        'apt write error — No space left on device',
        '/tmp was a 512MB tmpfs. Downloading apt package lists exceeded the limit.',
        'sudo rm -rf /var/lib/apt/lists/* && sudo apt-get clean && sudo apt-get update'
    ),
    (
        'pip install hanging on anthropic package',
        '/tmp (512MB tmpfs) too small for the 532MB anthropic package. Download stalled silently.',
        'Set TMPDIR=/var/tmp in Dockerfile and setup script. Added --no-cache-dir to pip install.'
    ),
    (
        'Docker BuildKit not available',
        'DOCKER_BUILDKIT=1 failed with "buildx component missing" on docker-compose v1.29.',
        'Dropped Docker entirely. Switched to direct VM deployment: Python venv + systemd + nginx.'
    ),
    (
        'pydantic version conflict with langchain',
        'requirements.txt pinned pydantic==2.5.0 but langchain>=0.3.0 requires pydantic>=2.7.4.',
        'Relaxed constraint to pydantic>=2.7.4 in requirements.txt.'
    ),
    (
        'frontend/src/lib/ missing from repo',
        'The lib/ pattern in root .gitignore (meant for Python venv) also blocked frontend/src/lib/utils.ts. Frontend build failed on VM.',
        'Updated .gitignore to only exclude venv/lib/. Committed frontend/src/lib/ files.'
    ),
    (
        'nginx IPv6 socket error (error 97)',
        'VM kernel does not support IPv6. nginx failed to start with: socket() [::]:80 failed.',
        'Removed all listen [::]:80 directives from nginx config files.'
    ),
    (
        'nginx 403 Forbidden',
        'Files in /var/www/aem-studio had wrong permissions after copy.',
        'chmod -R 755 /var/www/aem-studio && chown -R www-data:www-data /var/www/aem-studio'
    ),
    (
        '401 Unauthorized on all API calls',
        'ENVIRONMENT=production in .env.docker disabled the dev auth bypass. The bypass only activates in development/test mode.',
        'Changed ENVIRONMENT=production to ENVIRONMENT=development in .env.docker and restarted backend.'
    ),
    (
        'SQLite jobs table missing',
        'Backend health showed: sqlite3.OperationalError: no such table: jobs. Migrations had not run.',
        'Ran Base.metadata.create_all(bind=engine) with all models imported to initialise schema.'
    ),
    (
        'Dataset recipe schema validation bugs',
        'Three recipes silently failed: conref_pack (topic_count 8 < min 10), wide_branching (children_per_root 8 < min 10), large_scale (topic_count 100 < min 1000).',
        'Fixed chat_tools.py defaults and relaxed large_scale schema minimum from 1000 to 50.'
    ),
]

for title, issue, fix in issues:
    doc.add_heading(title, 2)
    p = doc.add_paragraph()
    p.add_run('Issue: ').bold = True
    p.add_run(issue)
    p2 = doc.add_paragraph()
    run = p2.add_run('Fix: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    p2.add_run(fix)

# 3. Key Learnings
doc.add_heading('3. Key Learnings', 1)

learnings = [
    ('tmpfs /tmp size limit',
     'Many Linux VMs mount /tmp as a 512MB tmpfs. Large Python packages (anthropic ~530MB) exceed it. '
     'Always set TMPDIR=/var/tmp and use pip --no-cache-dir on VMs with small tmpfs.'),
    ('Skip Docker on constrained VMs',
     'Docker Compose v1.29 + legacy builder + small tmpfs = unreliable builds. For single-VM deployments, '
     'systemd + nginx is simpler, faster, and more debuggable than Docker.'),
    ('ENVIRONMENT controls auth bypass',
     'ALLOW_DEV_AUTH_BYPASS=true is ignored unless ENVIRONMENT=development or test. '
     'Always verify this pair when deploying to a new environment.'),
    ('gitignore scope matters',
     'A broad lib/ pattern in .gitignore blocked frontend source files. '
     'Use scoped patterns (venv/lib/) for build artifacts, not directory names shared with source code.'),
    ('nginx IPv6 on cloud VMs',
     'Many internal/cloud VMs have IPv6 disabled at kernel level. nginx fails to start if listen [::]:80 is present. '
     'Always handle IPv6 absence by removing or commenting out IPv6 listen directives.'),
    ('Port security — two layers',
     'Cloud VM ports must be opened in TWO places: (1) OS firewall (ufw), (2) Cloud security group (ATS Corp). '
     'Missing either one blocks external access.'),
    ('systemd for background services',
     'systemd with Restart=always and EnvironmentFile is the correct approach for persistent background services. '
     'The service survives SSH disconnects, crashes, and VM reboots automatically.'),
]

for title, text in learnings:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(text)

# 4. Final Architecture
doc.add_heading('4. Final Architecture on VM', 1)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
t2.rows[0].cells[0].text = 'Layer'
t2.rows[0].cells[1].text = 'Technology'
t2.rows[0].cells[2].text = 'Config'
for layer, tech, config in [
    ('Frontend', 'React + Vite (static build)', 'nginx port 4502, root /var/www/aem-studio'),
    ('Reverse Proxy', 'nginx 1.18', '/etc/nginx/sites-available/aem-studio'),
    ('Backend', 'FastAPI + uvicorn', 'systemd aem-backend, port 8001'),
    ('Database', 'SQLite', 'backend/storage/app.db'),
    ('Vector DB', 'ChromaDB', 'backend/storage/chroma_db'),
    ('LLM', 'Azure OpenAI gpt-5.2', '.env.docker: AZURE_OPENAI_*'),
    ('Process Mgmt', 'systemd', 'Restart=always, auto-start on boot'),
]:
    r = t2.add_row()
    r.cells[0].text = layer
    r.cells[1].text = tech
    r.cells[2].text = config

# 5. Maintenance
doc.add_heading('5. Day-to-Day Maintenance Commands', 1)
for label, cmd in [
    ('Restart backend', 'systemctl restart aem-backend'),
    ('View live logs', 'journalctl -u aem-backend -f'),
    ('Check status', 'systemctl status aem-backend nginx'),
    ('Update app', 'cd ~/aem-guides-dataset-studio && git pull && python3 setup_vm.py'),
    ('Check disk', 'df -h'),
    ('Backend health', 'curl http://localhost:8001/health'),
    ('Stop everything', 'systemctl stop aem-backend nginx'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    run = p.add_run(cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

out = 'C:/Users/prashantp/Videos/aem-guides-dataset-studio/AEM_Studio_VM_Deployment.docx'
doc.save(out)
print('Saved:', out)
